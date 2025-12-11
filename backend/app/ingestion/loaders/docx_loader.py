"""
Microsoft Word document loader using python-docx.

Extracts text content from DOCX files while preserving paragraph structure.
"""

import logging
from typing import Dict, Any, Tuple
from pathlib import Path
import asyncio

try:
    from docx import Document
except ImportError:
    raise ImportError(
        "DOCX support requires python-docx. Install with: pip install python-docx"
    )

logger = logging.getLogger(__name__)


class DOCXLoader:
    """
    Loader for Microsoft Word documents (DOCX format).

    Extracts text content from Word documents, preserving paragraph structure
    and extracting document properties as metadata.
    """

    def __init__(self):
        """Initialize DOCX loader."""
        self.supported_extensions = [".docx"]

    async def load(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Load and extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Tuple of (extracted_text, metadata)

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is not a valid DOCX
        """
        try:
            file_path_obj = Path(file_path)

            if not file_path_obj.exists():
                raise FileNotFoundError(f"DOCX file not found: {file_path}")

            if file_path_obj.suffix.lower() not in self.supported_extensions:
                raise ValueError(f"File is not a DOCX: {file_path}")

            logger.info(f"Loading DOCX: {file_path}")

            # Run DOCX parsing in thread pool to avoid blocking
            text, metadata = await asyncio.to_thread(self._extract_text, file_path_obj)

            logger.info(
                f"Loaded DOCX: {metadata['paragraph_count']} paragraphs, "
                f"{len(text)} characters"
            )

            return text, metadata

        except Exception as e:
            logger.error(f"Failed to load DOCX {file_path}: {e}")
            raise

    def _extract_text(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from DOCX file (synchronous).

        Args:
            file_path: Path object for the DOCX file

        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            doc = Document(str(file_path))

            # Extract metadata from document properties
            metadata = {
                "source_type": "docx",
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
            }

            # Add document core properties if available
            try:
                core_props = doc.core_properties
                metadata["title"] = core_props.title or ""
                metadata["author"] = core_props.author or ""
                metadata["subject"] = core_props.subject or ""
                metadata["keywords"] = core_props.keywords or ""
                metadata["created"] = (
                    core_props.created.isoformat() if core_props.created else None
                )
                metadata["modified"] = (
                    core_props.modified.isoformat() if core_props.modified else None
                )
            except Exception as e:
                logger.debug(f"Could not extract document properties: {e}")

            # Extract text from paragraphs
            paragraphs = []
            all_text = []

            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    paragraphs.append({
                        "index": i,
                        "text": text,
                        "char_count": len(text),
                    })
                    all_text.append(text)

            # Combine paragraphs with double newline
            full_text = "\n\n".join(all_text)

            # Extract text from tables
            table_texts = []
            for table_num, table in enumerate(doc.tables, start=1):
                try:
                    table_data = self._extract_table_text(table, table_num)
                    if table_data:
                        table_texts.append(table_data)
                        # Add table to main text
                        full_text += f"\n\n{table_data}\n"
                except Exception as e:
                    logger.warning(f"Failed to extract table {table_num}: {e}")

            # Update metadata
            metadata["paragraph_count"] = len(paragraphs)
            metadata["table_count"] = len(table_texts)
            metadata["total_chars"] = len(full_text)
            metadata["paragraphs"] = paragraphs

            if not full_text.strip():
                logger.warning(f"No text extracted from DOCX: {file_path}")

            return full_text, metadata

        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise

    def _extract_table_text(self, table, table_num: int) -> str:
        """
        Extract text from a table in structured format.

        Args:
            table: python-docx Table object
            table_num: Table number for labeling

        Returns:
            Formatted table text
        """
        try:
            table_lines = [f"--- Table {table_num} ---"]

            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                # Join cells with pipe separator for table structure
                row_text = " | ".join(cells)
                if row_text.strip():
                    table_lines.append(row_text)

            return "\n".join(table_lines)

        except Exception as e:
            logger.warning(f"Error extracting table {table_num}: {e}")
            return ""

    async def load_with_formatting(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Load DOCX with additional formatting information.

        Extracts text along with formatting metadata like bold, italic, headings.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Tuple of (extracted_text, metadata_with_formatting)
        """
        try:
            file_path_obj = Path(file_path)

            if not file_path_obj.exists():
                raise FileNotFoundError(f"DOCX file not found: {file_path}")

            text, metadata = await asyncio.to_thread(
                self._extract_with_formatting, file_path_obj
            )

            return text, metadata

        except Exception as e:
            logger.error(f"Failed to load DOCX with formatting: {e}")
            raise

    def _extract_with_formatting(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text with formatting information (synchronous).

        Args:
            file_path: Path object for the DOCX file

        Returns:
            Tuple of (formatted_text, metadata)
        """
        try:
            doc = Document(str(file_path))

            metadata = {
                "source_type": "docx",
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
            }

            formatted_parts = []
            headings = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Check if paragraph is a heading
                if para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    headings.append({
                        "level": level,
                        "text": text,
                    })
                    # Add markdown-style heading
                    formatted_parts.append(f"\n\n{'#' * int(level)} {text}\n")
                else:
                    formatted_parts.append(text)

            full_text = "\n\n".join(formatted_parts)

            metadata["headings"] = headings
            metadata["heading_count"] = len(headings)
            metadata["total_chars"] = len(full_text)

            return full_text, metadata

        except Exception as e:
            logger.error(f"Error extracting formatted text: {e}")
            raise
